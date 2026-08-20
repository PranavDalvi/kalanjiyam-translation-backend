import io
import os
import tempfile
import pytest
from docx import Document
from fastapi.testclient import TestClient

import app.main as main
from app.main import chunk_for_indictrans

client = TestClient(main.app)


class MockTokenizer:
    """Mock tokenizer where 1 word ~ 1 token for predictable testing."""
    def __init__(self, token_ratio=1):
        self.token_ratio = token_ratio

    def encode(self, text: str, add_special_tokens: bool = False):
        if not text:
            return []
        words = text.split()
        if len(words) <= 1 and len(text) > 10:
            count = max(1, len(text) // 3)
        else:
            count = len(words)
        return [1] * (count * self.token_ratio)


# ============================================================================
# Case 1: Sentence split across physical lines
# ============================================================================
def test_sentence_split_across_physical_lines():
    input_text = "I hope to be able to make\nover charge during the week"
    tokenizer = MockTokenizer()
    chunks = chunk_for_indictrans(input_text, tokenizer=tokenizer, max_tokens=256)

    # Must be joined into a single logical translation chunk
    assert len(chunks) == 1
    assert chunks[0] == "I hope to be able to make over charge during the week"


# ============================================================================
# Case 2: Multiple paragraphs
# ============================================================================
def test_multiple_paragraphs_preserved():
    input_text = (
        "First paragraph line one\ncontinues here\n\n"
        "Second paragraph line one\ncontinues here"
    )
    tokenizer = MockTokenizer()
    chunks = chunk_for_indictrans(input_text, tokenizer=tokenizer, max_tokens=256)

    assert len(chunks) == 2
    assert chunks[0] == "First paragraph line one continues here"
    assert chunks[1] == "Second paragraph line one continues here"


# ============================================================================
# Case 3: Long paragraph exceeding max_tokens
# ============================================================================
def test_long_paragraph_split_at_sentence_boundaries():
    sentences = [
        "This is sentence number one.",
        "This is sentence number two.",
        "This is sentence number three.",
        "This is sentence number four.",
        "This is sentence number five.",
    ]
    # Each sentence is 5 words/tokens
    input_text = " ".join(sentences)
    tokenizer = MockTokenizer()
    # With max_tokens = 11, each chunk can hold at most 2 sentences (10 tokens)
    chunks = chunk_for_indictrans(input_text, tokenizer=tokenizer, max_tokens=11)

    assert len(chunks) == 3
    assert chunks[0] == "This is sentence number one. This is sentence number two."
    assert chunks[1] == "This is sentence number three. This is sentence number four."
    assert chunks[2] == "This is sentence number five."


# ============================================================================
# Case 4: Multiple PDF pages remain separate
# ============================================================================
def test_pdf_pages_remain_separate(monkeypatch):
    captured_translation_calls = []

    def fake_get_translation_model(model_name: str, src_lang_name: str, tgt_lang_name: str, gpu_id: int):
        return object(), MockTokenizer(), object()

    def fake_translate_batch_memory_safe(sentences, model, tokenizer, ip, src_lang, tgt_lang, batch_size, **kwargs):
        captured_translation_calls.append(list(sentences))
        return [f"TRANSLATED: {s}" for s in sentences]

    monkeypatch.setattr(main.service, "get_translation_model", fake_get_translation_model)
    monkeypatch.setattr(main.service, "translate_batch_memory_safe", fake_translate_batch_memory_safe)

    class MockPage:
        def __init__(self, text):
            self._text = text
        def extract_text(self):
            return self._text

    class MockPdf:
        def __init__(self, pages):
            self.pages = [MockPage(p) for p in pages]
        def __enter__(self):
            return self
        def __exit__(self, *args):
            pass

    import pdfplumber
    monkeypatch.setattr(pdfplumber, "open", lambda *args, **kwargs: MockPdf([
        "Page 1 first line\nPage 1 second line",
        "Page 2 first line\nPage 2 second line"
    ]))

    # Call /translate/document with PDF
    fake_pdf_bytes = b"%PDF-1.4 dummy"
    response = client.post(
        "/translate/document",
        data={
            "model_name": "ai4bharat/indictrans2-en-indic-1B",
            "source_language": "English",
            "target_language": "Hindi",
            "gpu_id": 0,
            "batch_size": 8,
        },
        files={"file": ("sample.pdf", fake_pdf_bytes, "application/pdf")},
    )

    assert response.status_code == 200
    # Must have 2 separate translation calls for the 2 pages
    assert len(captured_translation_calls) == 2
    assert captured_translation_calls[0] == ["Page 1 first line Page 1 second line"]
    assert captured_translation_calls[1] == ["Page 2 first line Page 2 second line"]


# ============================================================================
# Case 5: Empty lines / whitespace handling
# ============================================================================
def test_empty_lines_and_whitespace():
    tokenizer = MockTokenizer()
    assert chunk_for_indictrans("", tokenizer=tokenizer) == []
    assert chunk_for_indictrans("   \n\n\t  \n  ", tokenizer=tokenizer) == []

    # Paragraphs with extra blank lines in between
    input_text = "Para 1 line 1\n\n\n\n\nPara 2 line 1\n\n"
    chunks = chunk_for_indictrans(input_text, tokenizer=tokenizer)
    assert len(chunks) == 2
    assert chunks[0] == "Para 1 line 1"
    assert chunks[1] == "Para 2 line 1"


# ============================================================================
# Case 6: Very long single line
# ============================================================================
def test_very_long_single_line():
    # Long line with words
    long_line = " ".join([f"word{i}" for i in range(100)])
    tokenizer = MockTokenizer()
    chunks = chunk_for_indictrans(long_line, tokenizer=tokenizer, max_tokens=25)

    assert len(chunks) == 4
    for chunk in chunks:
        assert len(chunk.split()) <= 25

    # Extremely long single token without spaces
    giant_token = "A" * 300
    sliced_chunks = chunk_for_indictrans(giant_token, tokenizer=tokenizer, max_tokens=10)
    assert len(sliced_chunks) > 1
    assert "".join(sliced_chunks) == giant_token


# ============================================================================
# Case 7: Existing Gemma document translation path behaves as before
# ============================================================================
def test_gemma_document_translation_path_preserved(monkeypatch):
    called_proxy = []

    def fake_proxy_gemma(text, src_name, tgt_name, gpu_id=0, batch_size=4):
        called_proxy.append((text, src_name, tgt_name))
        return "\n".join([f"GEMMA_TRANS: {l}" for l in text.split("\n")])

    monkeypatch.setattr(main, "_proxy_gemma_translation", fake_proxy_gemma)

    # Test TXT document with Gemma model
    txt_content = b"Line 1\nLine 2\nLine 3"
    response = client.post(
        "/translate/document",
        data={
            "model_name": "google/gemma-4-12b-it",
            "source_language": "English",
            "target_language": "Tamil",
            "gpu_id": 0,
            "batch_size": 8,
        },
        files={"file": ("test.txt", txt_content, "text/plain")},
    )

    assert response.status_code == 200
    assert len(called_proxy) == 1
    assert called_proxy[0][1] == "English"
    assert called_proxy[0][2] == "Tamil"


# ============================================================================
# Case 8: Glossary / DNT handling preservation
# ============================================================================
def test_glossary_dnt_handling_with_chunking(monkeypatch):
    captured_input_sentences = []

    def fake_get_translation_model(model_name: str, src_lang_name: str, tgt_lang_name: str, gpu_id: int):
        return object(), MockTokenizer(), object()

    def fake_translate_batch_memory_safe(sentences, model, tokenizer, ip, src_lang, tgt_lang, batch_size, glossary_dict=None, **kwargs):
        captured_input_sentences.extend(sentences)
        # Verify glossary_dict was passed through intact
        assert glossary_dict == {"computer": "கணினி"}
        return [f"மொழிபெயர்க்கப்பட்ட: {s}" for s in sentences]

    monkeypatch.setattr(main.service, "get_translation_model", fake_get_translation_model)
    monkeypatch.setattr(main.service, "translate_batch_memory_safe", fake_translate_batch_memory_safe)
    monkeypatch.setattr(main.glossary_service, "get_merged_glossary_dict", lambda *args: {"computer": "கணினி"})

    txt_content = "The computer is an electronic device\nused everywhere today.\n\nComputers help people.".encode("utf-8")
    response = client.post(
        "/translate/document",
        data={
            "model_name": "ai4bharat/indictrans2-en-indic-1B",
            "source_language": "English",
            "target_language": "Tamil",
            "glossary": "comp",
            "gpu_id": 0,
            "batch_size": 8,
        },
        files={"file": ("doc.txt", txt_content, "text/plain")},
    )

    assert response.status_code == 200
    assert len(captured_input_sentences) == 2
    assert captured_input_sentences[0] == "The computer is an electronic device used everywhere today."
    assert captured_input_sentences[1] == "Computers help people."


# ============================================================================
# DOCX Process integration test with IndicTrans2
# ============================================================================
def test_docx_indictrans_process(monkeypatch):
    captured_batch = []

    def fake_translate_batch_memory_safe(sentences, model, tokenizer, ip, src_lang, tgt_lang, batch_size, **kwargs):
        captured_batch.extend(sentences)
        return [f"TRANS: {s}" for s in sentences]

    monkeypatch.setattr(main.service, "translate_batch_memory_safe", fake_translate_batch_memory_safe)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        docx_path = tf.name

    try:
        doc = Document()
        doc.add_paragraph("First sentence across\nphysical line break.")
        doc.add_paragraph("Second independent paragraph.")
        doc.save(docx_path)

        res_doc = main.service.process_docx(
            docx_path,
            model=object(),
            tokenizer=MockTokenizer(),
            ip=object(),
            src_lang="English",
            tgt_lang="Hindi",
            batch_size=8,
        )

        assert len(captured_batch) == 2
        assert captured_batch[0] == "First sentence across physical line break."
        assert captured_batch[1] == "Second independent paragraph."
        assert res_doc.paragraphs[0].text == "TRANS: First sentence across physical line break."
        assert res_doc.paragraphs[1].text == "TRANS: Second independent paragraph."
    finally:
        if os.path.exists(docx_path):
            os.remove(docx_path)
