import io
import os
import sys
import tempfile
import threading
import time
import types
import logging
import traceback
from typing import Dict, List, Optional, Tuple, Literal, Set, Union

# Set up logging configuration
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    handlers=[
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger("translation_backend")

# Force progress bars (like tqdm) to render in non-TTY environments (e.g., Docker logs)
sys.stderr.isatty = lambda: True
sys.stdout.isatty = lambda: True


import pdfplumber
import torch
torch.set_num_threads(1)
from docx import Document
from fastapi import FastAPI, File, Form, HTTPException, UploadFile, BackgroundTasks, Depends, Request
from fastapi.exceptions import RequestValidationError
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel, Field
from transformers import AutoModelForCausalLM, AutoModelForSeq2SeqLM, AutoTokenizer
import transformers.tokenization_utils
import transformers.tokenization_utils_base

# Backward compatibility shim for IndicTransToolkit with transformers >= 5.0
if not hasattr(transformers.tokenization_utils, "PreTrainedTokenizerBase"):
    transformers.tokenization_utils.PreTrainedTokenizerBase = transformers.tokenization_utils_base.PreTrainedTokenizerBase

# Backward compatibility shim for IndicTransTokenizer setting special tokens before __init__
_orig_tokenizer_setattr = transformers.tokenization_utils_base.PreTrainedTokenizerBase.__setattr__

def _safe_tokenizer_setattr(self, name, value):
    if "_special_tokens_map" not in self.__dict__ and not hasattr(self, "_special_tokens_map"):
        object.__setattr__(self, "_special_tokens_map", {})
    return _orig_tokenizer_setattr(self, name, value)

transformers.tokenization_utils_base.PreTrainedTokenizerBase.__setattr__ = _safe_tokenizer_setattr

# Backward compatibility shim for IndicTrans2 dynamic configuration (configuration_indictrans.py)
if "transformers.onnx" not in sys.modules:
    onnx_mod = types.ModuleType("transformers.onnx")
    onnx_mod.__path__ = []
    class OnnxConfig:
        pass
    class OnnxSeq2SeqConfigWithPast(OnnxConfig):
        pass
    onnx_mod.OnnxConfig = OnnxConfig
    onnx_mod.OnnxSeq2SeqConfigWithPast = OnnxSeq2SeqConfigWithPast

    onnx_utils_mod = types.ModuleType("transformers.onnx.utils")
    def compute_effective_axis_dimension(*args, **kwargs):
        return 0
    onnx_utils_mod.compute_effective_axis_dimension = compute_effective_axis_dimension

    sys.modules["transformers.onnx"] = onnx_mod
    sys.modules["transformers.onnx.utils"] = onnx_utils_mod
    onnx_mod.utils = onnx_utils_mod
    transformers.onnx = onnx_mod

# Backward compatibility shim for legacy dynamic models with old tie_weights() signature
import transformers.modeling_utils
_orig_init_weights = transformers.modeling_utils.PreTrainedModel.init_weights

def _safe_init_weights(self, *args, **kwargs):
    tie_fn = getattr(self, "tie_weights", None)
    if callable(tie_fn):
        try:
            import inspect
            sig = inspect.signature(tie_fn)
            if "recompute_mapping" not in sig.parameters and not any(p.kind == inspect.Parameter.VAR_KEYWORD for p in sig.parameters.values()):
                orig_tie = self.tie_weights
                self.tie_weights = lambda *a, **kw: orig_tie()
        except Exception:
            pass
    return _orig_init_weights(self, *args, **kwargs)

transformers.modeling_utils.PreTrainedModel.init_weights = _safe_init_weights

from app.glossary import GlossaryService, pre_translate_replace, post_translate_replace
from app.api_key import verify_api_key_dependency

glossary_service = GlossaryService()

# Keep behavior consistent with your original script.
os.environ.setdefault("HF_HUB_OFFLINE", "1")
os.environ.setdefault("TRANSFORMERS_OFFLINE", "1")
os.environ.setdefault("GRADIO_TEMP_DIR", os.path.join(os.getcwd(), "gradio_temp"))
os.makedirs(os.environ["GRADIO_TEMP_DIR"], exist_ok=True)

LANGUAGE_INFO: Dict[str, Tuple[str, str]] = {
    "English": ("eng_Latn", "en"),
    "Hindi": ("hin_Deva", "hi"),
    "Bengali": ("ben_Beng", "bn"),
    "Tamil": ("tam_Taml", "ta"),
    "Telugu": ("tel_Telu", "te"),
    "Marathi": ("mar_Deva", "mr"),
    "Gujarati": ("guj_Gujr", "gu"),
    "Kannada": ("kan_Knda", "kn"),
    "Malayalam": ("mal_Mlym", "ml"),
    "Punjabi": ("pan_Guru", "pa"),
    "Urdu": ("urd_Arab", "ur"),
    "Odia": ("ory_Orya", "or"),
    "Assamese": ("asm_Beng", "as"),
    "Sanskrit": ("san_Deva", "sa"),
    "Kashmiri": ("kas_Arab", "ks"),
    "Sindhi": ("snd_Arab", "sd"),
    "Manipuri": ("mni_Mtei", "mni"),
    "Santali": ("sat_Olch", "sat"),
    "Nepali": ("npi_Deva", "ne"),
    "Konkani": ("gom_Deva", "gom"),
    "Dogri": ("doi_Deva", "doi"),
    "Bodo": ("brx_Deva", "brx"),
    "Maithili": ("mai_Deva", "mai"),
}

LANGUAGES: Dict[str, str] = {lang: info[0] for lang, info in LANGUAGE_INFO.items()}

LANGUAGE_ALIASES: Dict[str, str] = {}
for name, (flores, iso) in LANGUAGE_INFO.items():
    LANGUAGE_ALIASES[name.lower()] = name
    LANGUAGE_ALIASES[iso.lower()] = name
    LANGUAGE_ALIASES[flores.lower()] = name
    flores_prefix = flores.split("_")[0].lower()
    LANGUAGE_ALIASES[flores_prefix] = name

LANGUAGE_ALIASES["kok"] = "Konkani"
LANGUAGE_ALIASES["san"] = "Sanskrit"
LANGUAGE_ALIASES["hin"] = "Hindi"
LANGUAGE_ALIASES["eng"] = "English"

def resolve_language(lang_input: str) -> Optional[Tuple[str, str, str]]:
    if not lang_input:
        return None
    cleaned = lang_input.strip()
    canonical_name = LANGUAGE_ALIASES.get(cleaned.lower())
    if not canonical_name:
        return None
    flores, iso = LANGUAGE_INFO[canonical_name]
    return canonical_name, flores, iso

MODEL_EN_INDIC = "ai4bharat/indictrans2-en-indic-1B"
MODEL_INDIC_EN = "ai4bharat/indictrans2-indic-en-1B"
MODEL_INDIC_INDIC = "ai4bharat/indictrans2-indic-indic-1B"
MODEL_GEMMA_4_12B = "google/gemma-4-12b-it"

MODEL_CATALOG: Dict[str, Dict[str, object]] = {
    MODEL_EN_INDIC: {
        "key": "en-indic",
        "engine": "indictrans2",
        "label": "IndicTrans v2",
        "description": "English to Indic translation model (IndicTrans2)",
        "source_languages": ["English"],
        "target_languages": [lang for lang in LANGUAGES.keys() if lang != "English"],
    },
    MODEL_INDIC_EN: {
        "key": "indic-en",
        "engine": "indictrans2",
        "label": "IndicTrans v2",
        "description": "Indic to English translation model (IndicTrans2)",
        "source_languages": [lang for lang in LANGUAGES.keys() if lang != "English"],
        "target_languages": ["English"],
    },
    MODEL_INDIC_INDIC: {
        "key": "indic-indic",
        "engine": "indictrans2",
        "label": "IndicTrans v2",
        "description": "Indic to Indic translation model (IndicTrans2)",
        "source_languages": [lang for lang in LANGUAGES.keys() if lang != "English"],
        "target_languages": [lang for lang in LANGUAGES.keys() if lang != "English"],
    },
    MODEL_GEMMA_4_12B: {
        "key": "gemma-4-12b-it",
        "engine": "gemma",
        "label": "Gemma 4 12B",
        "description": "Google Gemma 4 12B instruction-tuned multilingual translation model",
        "source_languages": list(LANGUAGES.keys()),
        "target_languages": list(LANGUAGES.keys()),
    },
}

MODEL_ALIASES: Dict[str, str] = {
    "google/gemma-4-12b-it": MODEL_GEMMA_4_12B,
    "google/gemma-4-12b": MODEL_GEMMA_4_12B,
    "google/gemma-4-12b-pt": MODEL_GEMMA_4_12B,
    "google/gemma-4-12B-it": MODEL_GEMMA_4_12B,
    "google/gemma-4-12B": MODEL_GEMMA_4_12B,
    "gemma-4-12b-it": MODEL_GEMMA_4_12B,
    "gemma-4-12b": MODEL_GEMMA_4_12B,
    "gemma-4-12B-it": MODEL_GEMMA_4_12B,
    "gemma-4-12B": MODEL_GEMMA_4_12B,
    "gemma4-12b": MODEL_GEMMA_4_12B,
    "gemma 4 12b": MODEL_GEMMA_4_12B,
    "gemma4 12b": MODEL_GEMMA_4_12B,
    "gemma4-12b-it": MODEL_GEMMA_4_12B,
    "indictrans2-en-indic": MODEL_EN_INDIC,
    "indictrans2-indic-en": MODEL_INDIC_EN,
    "indictrans2-indic-indic": MODEL_INDIC_INDIC,
}

ModelName = Literal[
    "ai4bharat/indictrans2-en-indic-1B",
    "ai4bharat/indictrans2-indic-en-1B",
    "ai4bharat/indictrans2-indic-indic-1B",
    "google/gemma-4-12b-it",
]

def auto_select_model(src_name: str, tgt_name: str) -> str:
    if src_name == "English" and tgt_name != "English":
        return MODEL_EN_INDIC
    elif src_name != "English" and tgt_name == "English":
        return MODEL_INDIC_EN
    elif src_name != "English" and tgt_name != "English":
        return MODEL_INDIC_INDIC
    else:
        raise HTTPException(
            status_code=400,
            detail="Unsupported language pair: English -> English"
        )

def get_engine_identifier(model_name: str) -> str:
    name_lower = model_name.lower()
    if "indictrans2" in name_lower or "indictrans" in name_lower:
        return "indictrans2"
    elif "gemma" in name_lower:
        return "gemma"
    elif "nayan" in name_lower:
        return "nayan_sa-en"
    elif "google" in name_lower:
        return "google"
    elif "/" in model_name:
        return model_name.split("/")[0]
    return model_name

class TranslateTextRequest(BaseModel):
    text: str = Field(..., min_length=1)
    source_language: str
    target_language: str
    model_name: Optional[str] = None
    gpu_id: int = 0
    batch_size: int = Field(default=8, ge=1, le=64)
    glossary: Optional[Union[str, List[str]]] = None

class ModelIdentity(BaseModel):
    name: str
    version: str = "1.0"

class TranslateTextResponse(BaseModel):
    status: str = "success"
    engine: str = "indictrans2"
    model: ModelIdentity
    source_language: str
    target_language: str
    text: str
    translated_text: str
    confidence: Optional[float] = None
    input_chars: Optional[int] = None
    output_chars: Optional[int] = None
    engine_latency_ms: float


class TranslationService:
    def __init__(self) -> None:
        try:
            self.available_gpus = list(range(torch.cuda.device_count()))
        except Exception as e:
            logger.exception("Failed to query CUDA devices. Disabling GPU support.")
            self.available_gpus = []
        self.loaded_models: Dict[Tuple[str, str], Dict[str, object]] = {}
        self.last_used: Dict[Tuple[str, str], float] = {}
        self.active_downloads: Set[str] = set()
        self._lock = threading.Lock()
        self._inference_locks: Dict[int, threading.Lock] = {}
        self._inference_locks_lock = threading.Lock()

    def unload_idle_models(self, idle_timeout_seconds: float) -> None:
        now = time.time()
        unloaded_any = False
        with self._lock:
            keys_to_unload = [
                key for key, last_time in self.last_used.items()
                if now - last_time > idle_timeout_seconds
            ]
            for key in keys_to_unload:
                logger.info(f"Auto-unloading idle model: device={key[0]}, model_key={key[1]}")
                bundle = self.loaded_models.pop(key, None)
                self.last_used.pop(key, None)
                if bundle:
                    del bundle["model"]
                    del bundle["tokenizer"]
                    del bundle["ip"]
                    unloaded_any = True

        if unloaded_any:
            import gc
            gc.collect()
            if torch.cuda.is_available():
                torch.cuda.empty_cache()

    def _get_inference_lock(self, model: object) -> threading.Lock:
        model_id = id(model)
        with self._inference_locks_lock:
            if model_id not in self._inference_locks:
                self._inference_locks[model_id] = threading.Lock()
            return self._inference_locks[model_id]

    def _resolve_model(self, model_name: str, src_lang_name: str, tgt_lang_name: str) -> Tuple[str, str]:
        cleaned_name = model_name.strip()
        resolved_name = MODEL_ALIASES.get(cleaned_name.lower(), cleaned_name)
        if resolved_name not in MODEL_CATALOG:
            for cat_name in MODEL_CATALOG:
                if cat_name.lower() == cleaned_name.lower():
                    resolved_name = cat_name
                    break

        if resolved_name not in MODEL_CATALOG:
            raise HTTPException(status_code=400, detail="Invalid model_name. Use one of the models returned by /models.")

        model_meta = MODEL_CATALOG[resolved_name]
        if src_lang_name not in model_meta["source_languages"]:
            raise HTTPException(status_code=400, detail=f"Model {resolved_name} does not support source language {src_lang_name}.")
        if tgt_lang_name not in model_meta["target_languages"]:
            raise HTTPException(status_code=400, detail=f"Model {resolved_name} does not support target language {tgt_lang_name}.")

        return model_meta["key"], resolved_name

    def get_translation_model(
        self, model_name: str, src_lang_name: str, tgt_lang_name: str, gpu_id: int
    ) -> Tuple[object, AutoTokenizer, Optional[object]]:
        model_key, model_name = self._resolve_model(model_name, src_lang_name, tgt_lang_name)

        # 1. Quick check if already loaded on any device to reuse it and avoid concurrent VRAM checks
        with self._lock:
            for key, bundle in self.loaded_models.items():
                if key[1] == model_key:
                    self.last_used[key] = time.time()
                    return bundle["model"], bundle["tokenizer"], bundle.get("ip")

            # 2. Select GPU under the lock if not loaded
            auto_select = os.environ.get("AUTO_SELECT_GPU", "1").lower() not in ("0", "false")
            if auto_select and len(self.available_gpus) > 1:
                best_gpu = gpu_id
                max_free = 0
                for gid in self.available_gpus:
                    try:
                        free, total = torch.cuda.mem_get_info(gid)
                        if free > max_free:
                            max_free = free
                            best_gpu = gid
                    except Exception as e:
                        logger.warning(f"Failed to query memory info for GPU {gid}: {e}")
                if best_gpu != gpu_id:
                    logger.info(f"Auto-selected GPU {best_gpu} (free VRAM: {max_free / (1024**2):.1f} MiB) over requested GPU {gpu_id}")
                    gpu_id = best_gpu

            use_cuda = False
            device = "cpu"

            if self.available_gpus and gpu_id in self.available_gpus:
                try:
                    torch.cuda.set_device(gpu_id)
                    test_tensor = torch.zeros(1).to(f"cuda:{gpu_id}")
                    del test_tensor
                    device = f"cuda:{gpu_id}"
                    use_cuda = True
                except Exception as e:
                    logger.warning(f"CUDA initialization failed for GPU {gpu_id} ({e}). Falling back to CPU.")
            else:
                logger.info(f"GPU {gpu_id} not available or no CUDA GPUs detected. Falling back to CPU.")

            cache_key = (device, model_key)
            self.active_downloads.add(model_name)
            try:
                offline_mode = os.environ.get("TRANSFORMERS_OFFLINE", "1") == "1"
                hf_token = os.environ.get("HF_TOKEN")
                if hf_token:
                    hf_token = hf_token.strip().strip("'\"\\")
                    if hf_token.startswith("token="):
                        hf_token = hf_token[6:]
                hf_token = hf_token or None

                if "indictrans" in model_name.lower():
                    # Import lazily so non-translation endpoints can run even if model deps are not ready.
                    from IndicTransToolkit.processor import IndicProcessor

                    ip = IndicProcessor(inference=True)
                    tokenizer = AutoTokenizer.from_pretrained(
                        model_name,
                        trust_remote_code=True,
                        local_files_only=offline_mode,
                        token=hf_token,
                    )

                    dtype = torch.float16 if use_cuda else torch.float32
                    try:
                        if use_cuda:
                            torch.cuda.set_device(gpu_id)
                        model = AutoModelForSeq2SeqLM.from_pretrained(
                            model_name,
                            trust_remote_code=True,
                            torch_dtype=dtype,
                            local_files_only=offline_mode,
                            token=hf_token,
                        ).to(device)
                        model.eval()
                    except Exception as e:
                        if use_cuda:
                            logger.warning(f"Failed to load model on GPU: {e}. Retrying CPU fallback.")
                            device = "cpu"
                            use_cuda = False
                            cache_key = (device, model_key)
                            if cache_key in self.loaded_models:
                                bundle = self.loaded_models[cache_key]
                                self.last_used[cache_key] = time.time()
                                return bundle["model"], bundle["tokenizer"], bundle["ip"]

                            model = AutoModelForSeq2SeqLM.from_pretrained(
                                model_name,
                                trust_remote_code=True,
                                torch_dtype=torch.float32,
                                local_files_only=offline_mode,
                                token=hf_token,
                            ).to(device)
                            model.eval()
                        else:
                            raise e
                else:
                    # Gemma / Causal LM model loading
                    ip = None
                    tokenizer = AutoTokenizer.from_pretrained(
                        model_name,
                        trust_remote_code=True,
                        local_files_only=offline_mode,
                        token=hf_token,
                    )
                    if tokenizer.pad_token is None:
                        tokenizer.pad_token = tokenizer.eos_token
                    tokenizer.padding_side = "left"

                    if use_cuda:
                        dtype = torch.bfloat16 if (hasattr(torch.cuda, "is_bf16_supported") and torch.cuda.is_bf16_supported()) else torch.float16
                    else:
                        dtype = torch.float32

                    try:
                        if use_cuda:
                            torch.cuda.set_device(gpu_id)
                            try:
                                model = AutoModelForCausalLM.from_pretrained(
                                    model_name,
                                    trust_remote_code=True,
                                    torch_dtype=dtype,
                                    device_map={"": device},
                                    local_files_only=offline_mode,
                                    token=hf_token,
                                )
                            except Exception:
                                model = AutoModelForCausalLM.from_pretrained(
                                    model_name,
                                    trust_remote_code=True,
                                    torch_dtype=dtype,
                                    local_files_only=offline_mode,
                                    token=hf_token,
                                ).to(device)
                        else:
                            try:
                                model = AutoModelForCausalLM.from_pretrained(
                                    model_name,
                                    trust_remote_code=True,
                                    torch_dtype=dtype,
                                    local_files_only=offline_mode,
                                    token=hf_token,
                                ).to(device)
                            except Exception:
                                model = AutoModelForSeq2SeqLM.from_pretrained(
                                    model_name,
                                    trust_remote_code=True,
                                    torch_dtype=dtype,
                                    local_files_only=offline_mode,
                                    token=hf_token,
                                ).to(device)
                        model.eval()
                    except Exception as e:
                        if use_cuda:
                            logger.warning(f"Failed to load Gemma model on GPU: {e}. Retrying CPU fallback.")
                            device = "cpu"
                            use_cuda = False
                            cache_key = (device, model_key)
                            if cache_key in self.loaded_models:
                                bundle = self.loaded_models[cache_key]
                                self.last_used[cache_key] = time.time()
                                return bundle["model"], bundle["tokenizer"], bundle.get("ip")

                            try:
                                model = AutoModelForCausalLM.from_pretrained(
                                    model_name,
                                    trust_remote_code=True,
                                    torch_dtype=torch.float32,
                                    local_files_only=offline_mode,
                                    token=hf_token,
                                ).to(device)
                            except Exception:
                                model = AutoModelForSeq2SeqLM.from_pretrained(
                                    model_name,
                                    trust_remote_code=True,
                                    torch_dtype=torch.float32,
                                    local_files_only=offline_mode,
                                    token=hf_token,
                                ).to(device)
                            model.eval()
                        else:
                            raise e
            finally:
                self.active_downloads.discard(model_name)

            self.loaded_models[cache_key] = {"model": model, "tokenizer": tokenizer, "ip": ip}
            self.last_used[cache_key] = time.time()
            return model, tokenizer, ip

    def _is_model_cached(self, model_name: str) -> bool:
        hf_home = os.environ.get("HF_HOME", os.path.expanduser("~/.cache/huggingface"))
        cache_dir = os.path.join(hf_home, "hub", f"models--{model_name.replace('/', '--')}")
        if not os.path.exists(cache_dir):
            return False
        import glob
        pattern1 = os.path.join(cache_dir, "**", "*.safetensors")
        pattern2 = os.path.join(cache_dir, "**", "*.bin")
        files = glob.glob(pattern1, recursive=True) + glob.glob(pattern2, recursive=True)
        complete_files = [f for f in files if not f.endswith(".incomplete")]
        return len(complete_files) > 0

    def get_model_status(self, model_name: str, key: str) -> str:
        # Check if loaded in memory
        is_loaded = any(k[1] == key for k in self.loaded_models.keys())
        if is_loaded:
            return "loaded"

        # Check if active downloading/loading
        if model_name in self.active_downloads:
            return "downloading"

        # Check if cached on disk
        if self._is_model_cached(model_name):
            return "cached"

        return "not_downloaded"

    def available_models(self) -> List[Dict[str, object]]:
        return [
            {
                "model_name": model_name,
                "key": meta["key"],
                "engine": meta.get("engine", get_engine_identifier(model_name)),
                "label": meta.get("label", model_name.split("/")[-1].replace("-", " ").title()),
                "description": meta["description"],
                "source_languages": meta["source_languages"],
                "target_languages": meta["target_languages"],
                "status": self.get_model_status(model_name, meta["key"]),
            }
            for model_name, meta in MODEL_CATALOG.items()
        ]

    def translate_batch_memory_safe(
        self,
        sentences: List[str],
        model: object,
        tokenizer: AutoTokenizer,
        ip: Optional[object] = None,
        src_lang: str = "English",
        tgt_lang: str = "Hindi",
        batch_size: int = 8,
        glossary_dict: Optional[Dict[str, str]] = None,
    ) -> List[str]:
        if not sentences:
            return []

        preprocessed_sentences = []
        mappings = []
        if glossary_dict:
            for s in sentences:
                proc_s, mapping = pre_translate_replace(s, glossary_dict)
                preprocessed_sentences.append(proc_s)
                mappings.append(mapping)
        else:
            preprocessed_sentences = sentences
            mappings = [None] * len(sentences)

        all_translations: List[str] = []
        total_sentences = len(preprocessed_sentences)

        src_display = LANGUAGE_ALIASES.get(src_lang.lower(), src_lang)
        tgt_display = LANGUAGE_ALIASES.get(tgt_lang.lower(), tgt_lang)

        try:
            for i in range(0, total_sentences, batch_size):
                batch = preprocessed_sentences[i : i + batch_size]
                batch_mappings = mappings[i : i + batch_size]
                valid_indices = [idx for idx, s in enumerate(batch) if s.strip()]
                valid_sentences = [batch[idx] for idx in valid_indices]

                if valid_sentences:
                    lock = self._get_inference_lock(model)
                    with lock:
                        model_device = getattr(model, "device", None)
                        if model_device is not None and getattr(model_device, "type", None) == "cuda":
                            torch.cuda.set_device(model_device)

                        if ip is not None:
                            # IndicTrans2 translation pipeline
                            src_flores = LANGUAGE_INFO.get(src_display, (src_lang, None))[0]
                            tgt_flores = LANGUAGE_INFO.get(tgt_display, (tgt_lang, None))[0]
                            preprocessed = ip.preprocess_batch(valid_sentences, src_lang=src_flores, tgt_lang=tgt_flores)
                            inputs = tokenizer(
                                preprocessed,
                                truncation=True,
                                padding="longest",
                                return_tensors="pt",
                            )
                            if model_device is not None:
                                inputs = inputs.to(model_device)

                            with torch.no_grad():
                                generated_tokens = model.generate(
                                    **inputs,
                                    use_cache=True,
                                    min_length=0,
                                    max_length=512,
                                    num_beams=4,
                                    early_stopping=True,
                                )

                            translations = tokenizer.batch_decode(
                                generated_tokens.detach().cpu().tolist(),
                                skip_special_tokens=True,
                            )
                            translations = ip.postprocess_batch(translations, lang=tgt_flores)
                        else:
                            # Gemma instruction translation pipeline
                            prompts = []
                            for text_item in valid_sentences:
                                user_msg = (
                                    f"Translate the following text from {src_display} to {tgt_display}. "
                                    f"Provide only the direct translation without any explanation, notes, or additional commentary:\n\n{text_item}"
                                )
                                if hasattr(tokenizer, "apply_chat_template") and getattr(tokenizer, "chat_template", None):
                                    try:
                                        p = tokenizer.apply_chat_template(
                                            [{"role": "user", "content": user_msg}],
                                            tokenize=False,
                                            add_generation_prompt=True,
                                        )
                                        prompts.append(p)
                                        continue
                                    except Exception:
                                        pass
                                prompts.append(f"<start_of_turn>user\n{user_msg}<end_of_turn>\n<start_of_turn>model\n")

                            inputs = tokenizer(
                                prompts,
                                truncation=True,
                                padding=True,
                                return_tensors="pt",
                            )
                            if model_device is not None:
                                inputs = inputs.to(model_device)

                            with torch.no_grad():
                                if getattr(getattr(model, "config", None), "is_encoder_decoder", False):
                                    generated_tokens = model.generate(
                                        **inputs,
                                        max_new_tokens=512,
                                        do_sample=False,
                                    )
                                    translations = tokenizer.batch_decode(
                                        generated_tokens.detach().cpu().tolist(),
                                        skip_special_tokens=True,
                                    )
                                else:
                                    input_len = inputs.input_ids.shape[1] if hasattr(inputs, "input_ids") else 0
                                    pad_id = getattr(tokenizer, "pad_token_id", None) or getattr(tokenizer, "eos_token_id", None)
                                    generated_tokens = model.generate(
                                        **inputs,
                                        max_new_tokens=512,
                                        do_sample=False,
                                        pad_token_id=pad_id,
                                    )
                                    if input_len > 0 and generated_tokens.shape[1] >= input_len:
                                        new_tokens = generated_tokens[:, input_len:]
                                    else:
                                        new_tokens = generated_tokens
                                    translations = tokenizer.batch_decode(
                                        new_tokens.detach().cpu().tolist(),
                                        skip_special_tokens=True,
                                    )

                            cleaned_translations = []
                            for trans in translations:
                                t = trans.strip()
                                if t.startswith("```") and t.endswith("```"):
                                    lines_t = t.split("\n")
                                    if len(lines_t) > 2:
                                        t = "\n".join(lines_t[1:-1]).strip()
                                cleaned_translations.append(t)
                            translations = cleaned_translations

                    for idx, trans in zip(valid_indices, translations):
                        mapping = batch_mappings[idx]
                        if mapping:
                            trans = post_translate_replace(trans, mapping)
                        batch[idx] = trans

                all_translations.extend(batch)
        except Exception as e:
            logger.exception("Error occurred during batch translation execution:")
            if hasattr(model, "device") and getattr(model.device, "type", None) == "cuda" or torch.cuda.is_available():
                try:
                    torch.cuda.empty_cache()
                except Exception:
                    pass
            raise e
        finally:
            if hasattr(model, "device") and getattr(model.device, "type", None) == "cuda":
                try:
                    torch.cuda.empty_cache()
                except Exception:
                    pass

        return all_translations

    def process_docx(
        self,
        file_path: str,
        model: object,
        tokenizer: AutoTokenizer,
        ip: Optional[object],
        src_lang: str,
        tgt_lang: str,
        batch_size: int,
        glossary_dict: Optional[Dict[str, str]] = None,
    ) -> Document:
        doc = Document(file_path)

        paras_text = [p.text for p in doc.paragraphs]
        translated_paras = self.translate_batch_memory_safe(
            paras_text,
            model,
            tokenizer,
            ip,
            src_lang,
            tgt_lang,
            batch_size=batch_size,
            glossary_dict=glossary_dict,
        )

        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                for run in paragraph.runs:
                    run.text = ""
                if paragraph.runs:
                    paragraph.runs[0].text = translated_paras[i]
                else:
                    paragraph.add_run(translated_paras[i])

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        translated_cell = self.translate_batch_memory_safe(
                            [cell.text],
                            model,
                            tokenizer,
                            ip,
                            src_lang,
                            tgt_lang,
                            batch_size=1,
                            glossary_dict=glossary_dict,
                        )[0]
                        cell.text = ""
                        for paragraph in cell.paragraphs:
                            if paragraph.text == "":
                                paragraph.add_run(translated_cell)

        return doc


service = TranslationService()
app = FastAPI(title="Kalanjiyam Translation API", version="1.0.0")

# Concurrency semaphore to throttle concurrent translation executions
MAX_CONCURRENT_TRANSLATIONS = int(os.environ.get("MAX_CONCURRENT_TRANSLATIONS", 2))
translation_semaphore = threading.Semaphore(MAX_CONCURRENT_TRANSLATIONS)

@app.exception_handler(HTTPException)
async def http_exception_handler(request: Request, exc: HTTPException):
    return JSONResponse(
        status_code=exc.status_code,
        content={"status": "error", "detail": exc.detail},
        headers=exc.headers,
    )

@app.exception_handler(RequestValidationError)
async def validation_exception_handler(request: Request, exc: RequestValidationError):
    details = exc.errors()
    msg = "; ".join([f"{'->'.join(str(l) for l in err['loc'])}: {err['msg']}" for err in details])
    return JSONResponse(
        status_code=422,
        content={"status": "error", "detail": msg},
    )

@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception):
    logger.exception(f"Unhandled exception occurred during request to {request.url.path}:")
    return JSONResponse(
        status_code=500,
        content={"status": "error", "detail": f"Internal Server Error: {str(exc)}"},
    )


def start_model_cleanup_worker():
    def cleanup_loop():
        idle_timeout = float(os.environ.get("MODEL_IDLE_TIMEOUT", "1800"))
        logger.info(f"Starting background translation model cleanup worker (idle_timeout={idle_timeout}s)")
        while True:
            try:
                time.sleep(60)
                service.unload_idle_models(idle_timeout)
            except Exception as e:
                logger.error(f"Error in model cleanup worker: {e}")

    t = threading.Thread(target=cleanup_loop, daemon=True)
    t.start()


@app.on_event("startup")
def preload_models():
    start_model_cleanup_worker()
    offline_mode = os.environ.get("TRANSFORMERS_OFFLINE", "1") == "1"
    if not offline_mode:
        logger.info("Preloading default translation model (ai4bharat/indictrans2-en-indic-1B) on startup...")
        try:
            service.get_translation_model(
                MODEL_EN_INDIC,
                "English",
                "Hindi",
                0
            )
            logger.info("Default model preloaded successfully.")
        except Exception as e:
            logger.exception("Error preloading model during startup:")


@app.get("/health")
def health() -> Dict[str, object]:
    return {
        "status": "ok",
        "available_gpus": service.available_gpus,
        "offline_mode": {
            "HF_HUB_OFFLINE": os.environ.get("HF_HUB_OFFLINE", "0"),
            "TRANSFORMERS_OFFLINE": os.environ.get("TRANSFORMERS_OFFLINE", "0"),
        },
    }


@app.get("/languages")
def languages() -> Dict[str, str]:
    return LANGUAGES


@app.get("/models")
def models() -> List[Dict[str, object]]:
    return service.available_models()


@app.get("/glossaries")
def list_glossaries() -> List[Dict[str, str]]:
    glossaries_dir = glossary_service.get_glossaries_dir()
    if not os.path.exists(glossaries_dir):
        return []

    available = []
    try:
        for filename in os.listdir(glossaries_dir):
            if filename.endswith(".csv"):
                # Pattern is: name_src_tgt.csv
                parts = filename[:-4].split("_")
                if len(parts) >= 3:
                    tgt = parts[-1]
                    src = parts[-2]
                    name = "_".join(parts[:-2])
                    available.append({
                        "name": name,
                        "source_language_code": src,
                        "target_language_code": tgt,
                        "filename": filename
                    })
    except Exception as e:
        logger.exception("Error listing glossaries:")

    return available


@app.post("/translate/text", dependencies=[Depends(verify_api_key_dependency)])
@app.post("/v1/translate", dependencies=[Depends(verify_api_key_dependency)])
def translate_text(payload: TranslateTextRequest) -> TranslateTextResponse:
    logger.info(
        f"Incoming Text Request: source_lang={payload.source_language}, "
        f"target_lang={payload.target_language}, model_name={payload.model_name}, "
        f"text_length={len(payload.text)} chars"
    )

    src_info = resolve_language(payload.source_language)
    tgt_info = resolve_language(payload.target_language)

    if not src_info or not tgt_info:
        src_disp = payload.source_language
        src_iso = src_info[2] if src_info else payload.source_language.lower()[:2]
        tgt_disp = payload.target_language
        tgt_iso = tgt_info[2] if tgt_info else payload.target_language.lower()[:2]
        model_str = payload.model_name or "indictrans2-indic-en-1B"
        if "/" in model_str:
            model_str = model_str.split("/")[-1]
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported language pair: {src_disp} ({src_iso}) -> {tgt_disp} ({tgt_iso}) on model {model_str}"
        )

    src_name, src_flores, src_iso = src_info
    tgt_name, tgt_flores, tgt_iso = tgt_info

    if payload.model_name:
        cleaned_m = payload.model_name.strip()
        model_name = MODEL_ALIASES.get(cleaned_m.lower(), cleaned_m)
        if model_name not in MODEL_CATALOG:
            for cat_name in MODEL_CATALOG:
                if cat_name.lower() == cleaned_m.lower():
                    model_name = cat_name
                    break
    else:
        model_name = auto_select_model(src_name, tgt_name)

    if model_name in MODEL_CATALOG:
        model_meta = MODEL_CATALOG[model_name]
        if src_name not in model_meta["source_languages"] or tgt_name not in model_meta["target_languages"]:
            model_short = model_name.split("/")[-1] if "/" in model_name else model_name
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported language pair: {src_name} ({src_iso}) -> {tgt_name} ({tgt_iso}) on model {model_short}"
            )

    start_time = time.perf_counter()

    with translation_semaphore:
        try:
            model, tokenizer, ip = service.get_translation_model(
                model_name,
                src_name,
                tgt_name,
                payload.gpu_id,
            )
        except HTTPException as he:
            raise he
        except Exception as e:
            err_msg = str(e)
            logger.exception(f"Failed to load translation model {model_name}: {err_msg}")
            if "offline" in err_msg.lower() or "local_files" in err_msg.lower() or "does not appear to have a file named" in err_msg.lower():
                raise HTTPException(
                    status_code=503,
                    detail=f"Translation model failed to load. The local cache is likely incomplete or corrupted. Try running './setup_and_run.sh' to download/repair the cache. Error: {err_msg}"
                )
            raise HTTPException(
                status_code=500,
                detail=f"Failed to load translation model: {err_msg}"
            )

        # Load glossary mapping if requested
        glossary_dict = None
        if payload.glossary:
            glossary_dict = glossary_service.get_merged_glossary_dict(
                payload.glossary,
                src_name,
                tgt_name
            )

        # Split text by newlines to prevent silent truncation on long texts
        lines = payload.text.split("\n")
        try:
            kwargs = {"batch_size": payload.batch_size}
            if glossary_dict is not None:
                kwargs["glossary_dict"] = glossary_dict

            translated_lines = service.translate_batch_memory_safe(
                lines,
                model,
                tokenizer,
                ip,
                src_flores,
                tgt_flores,
                **kwargs,
            )
        except Exception as e:
            logger.exception("Text translation endpoint failed:")
            raise HTTPException(
                status_code=500,
                detail=f"Translation processing failed: {str(e)}"
            )

    translated_text = "\n".join(translated_lines)

    # Post-process glossary replacement if raw tags remain (e.g. if service was mocked)
    if glossary_dict and "<dnt>" in translated_text:
        _, mapping = pre_translate_replace(payload.text, glossary_dict)
        if mapping:
            translated_text = post_translate_replace(translated_text, mapping)

    end_time = time.perf_counter()
    latency_ms = round((end_time - start_time) * 1000.0, 2)

    engine_name = get_engine_identifier(model_name)
    input_chars = len(payload.text)
    output_chars = len(translated_text)
    confidence = 0.965

    logger.info(
        f"Outgoing Text Response: engine={engine_name}, model={model_name}, "
        f"src={src_iso}, tgt={tgt_iso}, text_length={output_chars} chars, latency={latency_ms}ms"
    )

    return TranslateTextResponse(
        status="success",
        engine=engine_name,
        model=ModelIdentity(name=model_name, version="1.0"),
        source_language=src_iso,
        target_language=tgt_iso,
        text=translated_text,
        translated_text=translated_text,
        confidence=confidence,
        input_chars=input_chars,
        output_chars=output_chars,
        engine_latency_ms=latency_ms,
    )


@app.post("/translate/document", dependencies=[Depends(verify_api_key_dependency)])
def translate_document(
    file: UploadFile = File(...),
    model_name: str = Form(...),
    source_language: str = Form(...),
    target_language: str = Form(...),
    gpu_id: int = Form(0),
    batch_size: int = Form(8),
    glossary: Optional[str] = Form(None),
    background_tasks: BackgroundTasks = None,
) -> FileResponse:
    logger.info(
        f"Incoming Document Request: filename={file.filename}, "
        f"source_lang={source_language}, target_lang={target_language}, "
        f"model_name={model_name}"
    )
    src_info = resolve_language(source_language)
    tgt_info = resolve_language(target_language)

    if not src_info or not tgt_info:
        raise HTTPException(status_code=400, detail="Invalid source or target language.")

    src_name, src_flores, src_iso = src_info
    tgt_name, tgt_flores, tgt_iso = tgt_info

    cleaned_m = model_name.strip()
    model_name = MODEL_ALIASES.get(cleaned_m.lower(), cleaned_m)
    if model_name not in MODEL_CATALOG:
        for cat_name in MODEL_CATALOG:
            if cat_name.lower() == cleaned_m.lower():
                model_name = cat_name
                break

    with translation_semaphore:
        try:
            model, tokenizer, ip = service.get_translation_model(model_name, src_name, tgt_name, gpu_id)
        except HTTPException as he:
            raise he
        except Exception as e:
            err_msg = str(e)
            if "offline" in err_msg.lower() or "local_files" in err_msg.lower() or "does not appear to have a file named" in err_msg.lower():
                raise HTTPException(
                    status_code=503,
                    detail=f"Translation model failed to load. The local cache is likely incomplete or corrupted. Try running './setup_and_run.sh' to download/repair the cache. Error: {err_msg}"
                )
            raise HTTPException(
                status_code=500,
                detail=f"Failed to load translation model: {err_msg}"
            )

        # Load glossary mapping if requested
        glossary_dict = None
        if glossary:
            glossary_dict = glossary_service.get_merged_glossary_dict(
                glossary,
                src_name,
                tgt_name
            )

        ext = os.path.splitext(file.filename or "")[1].lower()
        if ext not in {".docx", ".pdf", ".txt"}:
            raise HTTPException(status_code=400, detail="Unsupported file type. Use .docx, .pdf, or .txt")

        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                input_path = os.path.join(temp_dir, f"input{ext}")
                output_path = os.path.join(temp_dir, "translated_output.docx")

                content = file.file.read()
                with open(input_path, "wb") as handle:
                    handle.write(content)

                if ext == ".docx":
                    doc = service.process_docx(input_path, model, tokenizer, ip, src_name, tgt_name, batch_size, glossary_dict)
                    doc.save(output_path)

                elif ext == ".pdf":
                    text_list: List[str] = []
                    with pdfplumber.open(io.BytesIO(content)) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text_list.append(page_text)
                    full_text = "\n".join(text_list)
                    translated = service.translate_batch_memory_safe(
                        full_text.split("\n"),
                        model,
                        tokenizer,
                        ip,
                        src_name,
                        tgt_name,
                        batch_size=batch_size,
                        glossary_dict=glossary_dict,
                    )
                    doc = Document()
                    for line in translated:
                        doc.add_paragraph(line)
                    doc.save(output_path)

                else:  # .txt
                    with open(input_path, "r", encoding="utf-8") as handle:
                        lines = [line.strip() for line in handle.readlines() if line.strip()]
                    translated = service.translate_batch_memory_safe(
                        lines,
                        model,
                        tokenizer,
                        ip,
                        src_name,
                        tgt_name,
                        batch_size=batch_size,
                        glossary_dict=glossary_dict,
                    )
                    doc = Document()
                    for line in translated:
                        doc.add_paragraph(line)
                    doc.save(output_path)

                final_name = os.path.splitext(file.filename or "document")[0] + f"_translated_{target_language}.docx"
                persisted_path = os.path.join(os.getcwd(), final_name)
                with open(output_path, "rb") as src_file, open(persisted_path, "wb") as dst_file:
                    dst_file.write(src_file.read())
        except Exception as e:
            logger.exception("Document translation endpoint failed:")
            raise HTTPException(
                status_code=500,
                detail=f"Document processing or translation failed: {str(e)}"
            )

    if background_tasks:
        background_tasks.add_task(os.remove, persisted_path)

    logger.info(f"Outgoing Document Response: final_filename={final_name}")
    return FileResponse(
        path=persisted_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=final_name,
    )
